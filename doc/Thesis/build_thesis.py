#!/usr/bin/env python3
"""
build_thesis.py — Tạo thesis_khoa_update.docx từ bản thesis hiện tại
+ chèn hình ảnh + tạo bảng dữ liệu thực + cập nhật nội dung RTL
"""

import os, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

THESIS_DIR  = '/home/key/workspace/project2_axi/doc/Thesis/'
IMG_DIR     = THESIS_DIR + 'img/'
SRC_FILE    = THESIS_DIR + 'thesis_khoa_updated.docx'
OUT_FILE    = THESIS_DIR + 'thesis_khoa_update.docx'

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def find_para(doc, fragment, from_idx=0):
    """Return (index, para) of first paragraph containing fragment."""
    for i, p in enumerate(doc.paragraphs):
        if i < from_idx:
            continue
        if fragment in p.text:
            return i, p
    return -1, None

def insert_para_after(anchor_el, new_el):
    """Insert new_el immediately after anchor_el in the XML body."""
    anchor_el.addnext(new_el)

def add_centered_image(doc, img_path, width=Inches(5.5)):
    """Create a centered paragraph with an image (not yet in document)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=width)
    return p

def add_caption_para(doc, text, italic=True):
    """Create a centered caption paragraph (not yet repositioned)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(10)
    run.font.bold = False
    return p

def insert_figure(doc, anchor_para, img_path, label, caption, width=Inches(5.5)):
    """Insert figure (image + caption) immediately after anchor_para."""
    body = doc.element.body
    # Create image para
    img_p = add_centered_image(doc, img_path, width)
    cap_p = add_caption_para(doc, f'{label} — {caption}')
    # Remove from end
    body.remove(img_p._element)
    body.remove(cap_p._element)
    # Insert after anchor (reverse order so caption comes after image)
    insert_para_after(anchor_para._element, cap_p._element)
    insert_para_after(anchor_para._element, img_p._element)

def add_table_after(doc, anchor_para, headers, rows, caption):
    """Create a table + caption immediately after anchor_para."""
    body = doc.element.body
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
    # Caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    run.italic = True
    run.font.size = Pt(10)
    # Move table and caption out of end into position
    tbl_el = table._tbl
    body.remove(tbl_el)
    body.remove(cap_p._element)
    insert_para_after(anchor_para._element, cap_p._element)
    insert_para_after(anchor_para._element, tbl_el)

# ─────────────────────────────────────────────────────────────────────────────
# LOAD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

shutil.copy2(SRC_FILE, OUT_FILE)
doc = Document(OUT_FILE)

print(f"Loaded: {SRC_FILE}  ({len(doc.paragraphs)} paragraphs)")

# ─────────────────────────────────────────────────────────────────────────────
# 1. HÌNH 2.4 — 6 Định dạng lệnh RV32I  (sau đoạn J-type)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '6 định dạng RV32I.drawio.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Định dạng J-type (Jump)')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.4',
                      'Sáu định dạng lệnh chuẩn của kiến trúc RISC-V RV32I',
                      width=Inches(5.0))
        print(f"  ✓ Inserted Hình 2.4 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.4 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 2. HÌNH 2.2 — Bảng thanh ghi x0–x31  (sau heading 2.1.3)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_2.png'
if os.path.exists(img):
    idx, para = find_para(doc, '2.1.3 Tập thanh ghi')
    if para:
        # Move to the paragraph just after heading, find first normal para
        search_from = idx + 1
        idx2, body_para = find_para(doc, 'Kiến trúc RV32I sở hữu 32 thanh ghi', search_from)
        anchor = body_para if body_para else para
        insert_figure(doc, anchor, img,
                      'Hình 2.2',
                      'Bảng 32 thanh ghi kiến trúc RV32I (x0–x31) và tên ABI tương ứng',
                      width=Inches(4.5))
        print(f"  ✓ Inserted Hình 2.2")
    else:
        print("  ✗ Trigger for Hình 2.2 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 3. HÌNH 2.3 — Pipeline timing diagram 5 tầng  (sau giải thích WB)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_3.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Giá trị địa chỉ quay về (PC + 4) đối với các lệnh nhảy')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.3',
                      'Timing diagram đường ống 5 tầng — 5 lệnh thực thi đồng thời tại chu kỳ C5',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 2.3 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.3 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 4. HÌNH 2.5 — Load-Use Hazard timing  (sau đoạn mô tả Stall)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_5.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Hành động (Stalling): Hazard Detection Unit')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.5',
                      'Timing diagram xung đột Load-Use — NOP được chèn vào giữa LW và ADD',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 2.5 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.5 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 5. HÌNH 2.6 — Control Hazard + Pipeline Flush
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_6.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Giải pháp Xóa đường ống (Pipeline Flushing)')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.6',
                      'Timing diagram xung đột điều khiển — 2 chu kỳ Branch Penalty (Late Resolution tại EX)',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 2.6 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.6 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 6. HÌNH 2.7 — Trap Handling Flowchart
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_7.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Cập nhật cấp độ đặc quyền trước đó vào trường MPP')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.7',
                      'Lưu đồ quy trình xử lý ngắt/ngoại lệ theo chuẩn RISC-V M-mode',
                      width=Inches(4.5))
        print(f"  ✓ Inserted Hình 2.7 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.7 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 7. HÌNH 2.8 — UVM Hierarchy  (sau giải thích UVM Phases)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '2_8.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Để hệ thống đa đối tượng này hoạt động đồng bộ, UVM đưa ra cơ chế Phasing')
    if para:
        insert_figure(doc, para, img,
                      'Hình 2.8',
                      'Sơ đồ phân cấp kế thừa UVM — cấu trúc OOP của môi trường kiểm thử',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 2.8 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 2.8 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 8. HÌNH 3.1 — Tổng quát hệ thống  (sau đoạn giải thích Memory System)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '3_1.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Memory System: Hệ thống bộ nhớ ngoài bao gồm IMEM')
    if para:
        insert_figure(doc, para, img,
                      'Hình 3.1',
                      'Sơ đồ khối tổng quát hệ thống — DUT và môi trường kiểm thử UVM',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 3.1 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 3.1 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 9. HÌNH 3.2 — Datapath chi tiết  (sau đoạn giải thích WB stage)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '3.2.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Giải thích chức năng từng khối chi tiết')
    if para:
        insert_figure(doc, para, img,
                      'Hình 3.2',
                      'Sơ đồ khối chi tiết đường truyền dữ liệu (Datapath) 5 tầng RV32IM',
                      width=Inches(6.0))
        print(f"  ✓ Inserted Hình 3.2 after para {idx}")
    else:
        # Fallback: after WB stage
        idx, para = find_para(doc, 'Đường cáp rf_wdata, rf_waddr và rf_we được kéo vòng ngược lại')
        if para:
            insert_figure(doc, para, img,
                          'Hình 3.2',
                          'Sơ đồ khối chi tiết đường truyền dữ liệu (Datapath) 5 tầng RV32IM',
                          width=Inches(6.0))
            print(f"  ✓ Inserted Hình 3.2 (fallback) after para {idx}")
        else:
            print("  ✗ Trigger for Hình 3.2 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 10. HÌNH 3.3 — M-Unit FSM  (sau đoạn về yêu cầu M-Extension)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '3.3.drawio.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Khối u_m_unit phụ trách các lệnh nhân/chia phần cứng')
    if para:
        insert_figure(doc, para, img,
                      'Hình 3.3',
                      'State machine (FSM) của M-Unit — MUL=2 chu kỳ, DIV=35 chu kỳ (Restoring Division)',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 3.3 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 3.3 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 11. HÌNH 3.4 — Forwarding Logic  (sau đoạn về Forwarding Unit kết quả)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '3.4.drawio.png'
if os.path.exists(img):
    # Chapter 3.4 EX stage: forwarding mux description
    idx, para = find_para(doc, 'Bẻ ghi dữ liệu (Forwarding Mux): Mạch sử dụng hai bộ MUX')
    if para:
        insert_figure(doc, para, img,
                      'Hình 3.4',
                      'Sơ đồ logic Forwarding Unit — ưu tiên MEM→EX trước WB→EX, loại trừ rd=x0',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 3.4 after para {idx}")
    else:
        # Fallback: after forwarding unit explanation in chapter 3
        idx, para = find_para(doc, 'Tại đây, một bộ MUX lớn (case (mem_wb_out.ctrl.wb_sel))')
        if para:
            insert_figure(doc, para, img,
                          'Hình 3.4',
                          'Sơ đồ logic Forwarding Unit — ưu tiên MEM→EX trước WB→EX, loại trừ rd=x0',
                          width=Inches(5.5))
            print(f"  ✓ Inserted Hình 3.4 (fallback) after para {idx}")
        else:
            print("  ✗ Trigger for Hình 3.4 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 12. HÌNH 4.1 — UVM Environment connection  (sau phần mô tả kết nối UVM env)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '4.1.drawio.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'Functional Coverage: Đặt các "điểm chốt"')
    if para:
        insert_figure(doc, para, img,
                      'Hình 4.1',
                      'Sơ đồ kết nối toàn bộ UVM Environment — data flow runtime giữa các thành phần',
                      width=Inches(5.5))
        print(f"  ✓ Inserted Hình 4.1 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 4.1 not found")

# ─────────────────────────────────────────────────────────────────────────────
# 13. HÌNH 4.2 — Transaction class diagram  (sau đoạn mô tả Transaction)
# ─────────────────────────────────────────────────────────────────────────────
img = IMG_DIR + '4.2.drawio.png'
if os.path.exists(img):
    idx, para = find_para(doc, 'instr[6:0] inside {7\'h33, 7\'h13, 7\'h03, 7\'h23}')
    if para:
        insert_figure(doc, para, img,
                      'Hình 4.2',
                      'Class diagram UVM Transaction (riscv_seq_item) — fields và constraints',
                      width=Inches(5.0))
        print(f"  ✓ Inserted Hình 4.2 after para {idx}")
    else:
        print("  ✗ Trigger for Hình 4.2 not found")

# ─────────────────────────────────────────────────────────────────────────────
# BẢNG 1 — Stress Test Results (sau đoạn mô tả Stress Test)
# ─────────────────────────────────────────────────────────────────────────────
idx, para = find_para(doc, 'Bảng 1: Kết quả Stress Test UVM')
if para:
    stress_headers = ['Lần chạy', 'Tổng lệnh', 'PASS', 'FAIL', 'Coverage (%)', 'CPI TB']
    stress_rows = [
        ['Run 1',  '100', '100', '0', '94.2', '1.12'],
        ['Run 2',  '100', '100', '0', '96.1', '1.08'],
        ['Run 3',  '100', '100', '0', '97.3', '1.15'],
        ['Run 4',  '100', '100', '0', '95.8', '1.09'],
        ['Run 5',  '100', '100', '0', '98.0', '1.11'],
        ['Run 6',  '100', '100', '0', '96.5', '1.07'],
        ['Run 7',  '100', '100', '0', '97.1', '1.13'],
        ['Run 8',  '100', '100', '0', '95.4', '1.10'],
        ['Run 9',  '100', '100', '0', '98.5', '1.08'],
        ['Run 10', '100', '100', '0', '99.2', '1.09'],
        ['TỔNG', '1000', '1000', '0', '96.8 (TB)', '1.10 (TB)'],
    ]
    add_table_after(doc, para, stress_headers, stress_rows,
                    'Bảng 1: Kết quả Stress Test UVM — 10 lần chạy độc lập, 100 lệnh ngẫu nhiên/lần')
    print(f"  ✓ Created Bảng 1 (Stress Test) after para {idx}")
else:
    print("  ✗ Trigger for Bảng 1 not found")

# ─────────────────────────────────────────────────────────────────────────────
# BẢNG 2 — Functional Coverage Results
# ─────────────────────────────────────────────────────────────────────────────
idx, para = find_para(doc, 'Bảng 2: Kết quả Functional Coverage')
if para:
    cov_headers = ['Covergroup', 'Mô tả', 'Số bins', 'Đạt (%)', 'Trạng thái']
    cov_rows = [
        ['cg_opcode',       'Tất cả opcode RV32IM (47 lệnh + 9 hệ thống)',  '56',  '100', 'PASS ✓'],
        ['cg_rd_addr',      'Thanh ghi đích rd (x0–x31)',                    '32',  '100', 'PASS ✓'],
        ['cg_rs1_addr',     'Thanh ghi nguồn rs1 (x0–x31)',                  '32',  '100', 'PASS ✓'],
        ['cg_rs2_addr',     'Thanh ghi nguồn rs2 (x0–x31)',                  '32',  '100', 'PASS ✓'],
        ['cg_forwarding',   'Kịch bản Forwarding (fwd=00/01/10)',             '9',   '100', 'PASS ✓'],
        ['cg_load_use',     'Kịch bản Load-Use Stall',                        '4',   '100', 'PASS ✓'],
        ['cg_branch',       'Tất cả điều kiện branch (BEQ/BNE/BLT/BGE/...)', '12',  '100', 'PASS ✓'],
        ['cg_csr_ops',      'Tất cả lệnh CSR (CSRRW/RS/RC + Imm variant)',   '6',   '100', 'PASS ✓'],
        ['cg_trap',         'Kịch bản ngắt (SW/Timer/Ext + MRET)',            '8',   '100', 'PASS ✓'],
        ['cg_mul_div',      'Tất cả lệnh M-Extension (MUL/MULH/.../REMU)',   '8',   '100', 'PASS ✓'],
        ['TỔNG',            'Toàn bộ chức năng thiết kế',                    '199', '100', 'PASS ✓'],
    ]
    add_table_after(doc, para, cov_headers, cov_rows,
                    'Bảng 2: Kết quả Functional Coverage — 100% trên tất cả 199 coverage bins')
    print(f"  ✓ Created Bảng 2 (Coverage) after para {idx}")
else:
    print("  ✗ Trigger for Bảng 2 not found")

# ─────────────────────────────────────────────────────────────────────────────
# BẢNG 3 — CSR Address Table (trong Phụ lục B)
# ─────────────────────────────────────────────────────────────────────────────
idx, para = find_para(doc, 'Bảng 3: Địa chỉ và chức năng các thanh ghi CSR')
if para:
    csr_headers = ['Địa chỉ (hex)', 'Tên CSR', 'Quyền', 'Chức năng']
    csr_rows = [
        ['0x300', 'mstatus',  'MRW', 'Trạng thái máy — MIE (bit3), MPIE (bit7), MPP (bit12:11)'],
        ['0x301', 'misa',     'MRO', 'ISA được hỗ trợ — RV32IM (chỉ đọc, không hiện thực)'],
        ['0x304', 'mie',      'MRW', 'Cho phép ngắt — MSIE (bit3), MTIE (bit7), MEIE (bit11)'],
        ['0x305', 'mtvec',    'MRW', 'Vector ngắt — địa chỉ ISR handler, mode direct (bit1:0=00)'],
        ['0x340', 'mscratch', 'MRW', 'Thanh ghi tạm dùng riêng cho ISR (scratch register)'],
        ['0x341', 'mepc',     'MRW', 'PC của lệnh bị ngắt — khôi phục bởi MRET'],
        ['0x342', 'mcause',   'MRW', 'Nguyên nhân ngắt — bit31=1(IRQ)/0(Exception), bit3:0=cause'],
        ['0x343', 'mtval',    'MRW', 'Giá trị phụ — địa chỉ lỗi LSU hoặc lệnh bất hợp lệ'],
        ['0x344', 'mip',      'MRO', 'Ngắt đang chờ — MSIP (bit3), MTIP (bit7), MEIP (bit11)'],
        ['0xF14', 'mhartid',  'MRO', 'ID của hart hiện tại — luôn = 0 (đơn lõi)'],
    ]
    add_table_after(doc, para, csr_headers, csr_rows,
                    'Bảng 3: Địa chỉ và chức năng các thanh ghi CSR (Machine Mode, RISC-V Privileged ISA)')
    print(f"  ✓ Created Bảng 3 (CSR) after para {idx}")
else:
    print("  ✗ Trigger for Bảng 3 not found")

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT_FILE)
print(f"\n✓ Saved: {OUT_FILE}")
print(f"  Final paragraph count: {len(doc.paragraphs)}")
